"""
DocxHandler 单元测试
=====================
所有 .docx blob 在测试中现场用 python-docx 构造到 BytesIO，
不依赖任何外部 fixture 文件。

# 人工编写
"""
import sys
import unittest
from io import BytesIO
from pathlib import Path

ROOT = Path(__file__).parent.parent / "code"
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from core.handlers import HandlerRegistry  # noqa: E402
from core.handlers.docx_handler import DocxHandler  # noqa: E402  触发注册


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """构造一个内存 .docx，返回字节流。"""
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for i, row in enumerate(table_rows):
            for j, cell in enumerate(row):
                table.rows[i].cells[j].text = cell
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtractText(unittest.TestCase):
    def setUp(self):
        self.h = DocxHandler()

    def test_extract_paragraphs(self):
        blob = make_docx(["第一段", "第二段", "第三段"])
        out = self.h.extract_text(blob)
        self.assertIsNotNone(out)
        lines = out.splitlines()
        self.assertEqual(lines[0], "第一段")
        self.assertEqual(lines[1], "第二段")
        self.assertEqual(lines[2], "第三段")

    def test_extract_preserves_empty_paragraphs(self):
        # 空段落必须保留为空行，否则结构 diff 看不出来
        blob = make_docx(["A", "", "B"])
        out = self.h.extract_text(blob)
        lines = out.splitlines()
        self.assertEqual(lines, ["A", "", "B"])

    def test_extract_includes_table_cells(self):
        blob = make_docx(["前言"], table_rows=[["姓名", "成绩"], ["张三", "90"]])
        out = self.h.extract_text(blob)
        self.assertIn("前言", out)
        self.assertIn("姓名", out)
        self.assertIn("成绩", out)
        self.assertIn("张三", out)
        self.assertIn("90", out)


class TestDescribeChange(unittest.TestCase):
    def setUp(self):
        self.h = DocxHandler()

    def test_describe_pure_add(self):
        old = make_docx(["A"])
        new = make_docx(["A", "B", "C"])
        desc = self.h.describe_change(old, new)
        self.assertIn("增加 2 段", desc)
        self.assertIn("删除 0 段", desc)
        self.assertIn("修改 0 段", desc)

    def test_describe_pure_remove(self):
        old = make_docx(["A", "B", "C"])
        new = make_docx(["A"])
        desc = self.h.describe_change(old, new)
        self.assertIn("增加 0 段", desc)
        self.assertIn("删除 2 段", desc)
        self.assertIn("修改 0 段", desc)

    def test_describe_modify(self):
        old = make_docx(["第一段", "原始内容", "结尾"])
        new = make_docx(["第一段", "修改后的内容", "结尾"])
        desc = self.h.describe_change(old, new)
        self.assertIn("修改 1 段", desc)

    def test_describe_no_content_change(self):
        # 同样段落、不同 docx 包（字节不同但段落一致）
        old = make_docx(["相同", "段落"])
        new = make_docx(["相同", "段落"])
        desc = self.h.describe_change(old, new)
        self.assertIn("一致", desc)


class TestRenderDiff(unittest.TestCase):
    def setUp(self):
        self.h = DocxHandler()

    def test_render_diff_tags(self):
        old = make_docx(["A", "B", "C"])
        new = make_docx(["A", "X", "C"])
        out = self.h.render_diff(old, new)
        tags = [t for t, _ in out]
        self.assertIn("added", tags)
        self.assertIn("removed", tags)
        self.assertIn("normal", tags)
        # 不会出现 '? ' hint
        self.assertNotIn("?", tags)
        contents = " ".join(c for _, c in out)
        self.assertIn("B", contents)
        self.assertIn("X", contents)

    def test_render_diff_table_change(self):
        old = make_docx(["头"], table_rows=[["姓名", "成绩"], ["张三", "90"]])
        new = make_docx(["头"], table_rows=[["姓名", "成绩"], ["张三", "95"]])
        out = self.h.render_diff(old, new)
        contents = [c for _, c in out]
        # 应能看到 90 被移除、95 被增加
        self.assertTrue(any("90" in c for c in contents))
        self.assertTrue(any("95" in c for c in contents))


class TestRegistry(unittest.TestCase):
    def test_docx_extension_registered(self):
        h = HandlerRegistry.for_path(Path("report.docx"))
        self.assertIsInstance(h, DocxHandler)

    def test_docx_extension_case_insensitive(self):
        h = HandlerRegistry.for_path(Path("Report.DOCX"))
        self.assertIsInstance(h, DocxHandler)


class TestRobustness(unittest.TestCase):
    def setUp(self):
        self.h = DocxHandler()

    def test_corrupt_blob_extract_returns_none(self):
        out = self.h.extract_text(b"this is not a real docx zip package")
        self.assertIsNone(out)

    def test_corrupt_blob_describe_returns_meta(self):
        # 异常不能冒到 UI；返回失败提示字符串
        desc = self.h.describe_change(b"garbage", b"more garbage")
        self.assertIn("失败", desc)

    def test_corrupt_blob_render_returns_meta(self):
        out = self.h.render_diff(b"garbage", b"more garbage")
        self.assertEqual(len(out), 1)
        self.assertEqual(out[0][0], "meta")
        self.assertIn("失败", out[0][1])


class TestHeaderFooterTracking(unittest.TestCase):
    """追踪盲区修复：页眉/页脚 + 文本框文字必须可见。"""

    def setUp(self):
        self.h = DocxHandler()

    @staticmethod
    def _docx_with_header(body: str, header: str) -> bytes:
        doc = Document()
        doc.add_paragraph(body)
        doc.sections[0].header.paragraphs[0].text = header
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_header_text_extracted(self):
        blob = self._docx_with_header("正文", "我的页眉")
        out = self.h.extract_text(blob)
        self.assertIn("我的页眉", out)
        self.assertIn("页眉", out)  # 带标记前缀

    def test_header_change_is_visible(self):
        # 正文不变、只改页眉——原版会误报"段落级一致"
        old = self._docx_with_header("正文不变", "页眉第一版")
        new = self._docx_with_header("正文不变", "页眉第二版")
        desc = self.h.describe_change(old, new)
        self.assertNotIn("一致", desc)
        contents = " ".join(c for _, c in self.h.render_diff(old, new))
        self.assertIn("页眉第一版", contents)
        self.assertIn("页眉第二版", contents)


if __name__ == "__main__":
    unittest.main(verbosity=2)
