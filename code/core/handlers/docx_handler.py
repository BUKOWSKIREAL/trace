"""
DocxHandler — .docx 文档的段落级 diff / 摘要
==============================================
通过 python-docx 把 .docx 解析成"段落 + 表格单元格"的纯文本流，
再走 difflib 做行级（实为段落级）比对。

# AI生成 (Claude 4.7 协助)
# 人工修正：
#   - 段落与表格单元格一律按出现顺序拼成一条文本流，空段落保留空行，
#     这样表格行被增删时 diff 才会显式地多/少几行；
#   - python-docx 解析坏包会抛 zipfile.BadZipFile / KeyError 等多种异常，
#     一律 catch 成 meta 行，避免异常冒到 UI 层。
#   - 追踪盲区修复：补齐页眉/页脚 + XML <w:t> 文本框文字——原版只读
#     doc.paragraphs+tables，改页眉/文本框会被误报"段落级一致"。
"""
from __future__ import annotations

import difflib
import logging
from io import BytesIO

from core.handlers.base import DiffLine, FileHandler, HandlerRegistry

logger = logging.getLogger("trace.handlers.docx")


class DocxHandler(FileHandler):
    """
    .docx 文件 handler，做段落级 diff。
    """

    extensions = [".docx"]

    def extract_text(self, blob: bytes) -> str | None:
        """
        把所有段落、表格单元格、页眉/页脚按顺序拼成一段文本，每行一个语义单元。
        损坏的 docx 返回 None——调用方应据此把它当成二进制处理。

        # 人工修正（追踪盲区修复）：原版只读 doc.paragraphs + doc.tables，
        # 漏掉了用户实际会编辑的三类文本——导致"改了文档却显示没变化"：
        #   1. 页眉 / 页脚（doc.sections[*].header/footer）
        #   2. 文本框 / SmartArt / 形状里的文字（不在 doc.paragraphs 里，
        #      但都是 XML 里的 <w:t> 文本节点）
        # 表格单元格本身就含段落，正文段落和单元格已覆盖；这里用一个全文档
        # <w:t> 兜底集合补齐"正文流之外"的文本，且去重避免和正文重复展示。
        """
        try:
            from docx import Document  # 延迟导入，避免无 docx 时整模块崩
        except ImportError:
            logger.warning("python-docx 未安装，DocxHandler 退化为不可解析")
            return None

        try:
            doc = Document(BytesIO(blob))
        except Exception as e:
            logger.warning("DocxHandler 解析失败：%s", e)
            return None

        lines: list[str] = []
        # 段落（空段落保留为空行，便于段落级 diff 可见结构差异）
        for p in doc.paragraphs:
            lines.append(p.text)
        # 表格——按行展开，单元格也作为独立段落参与 diff
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    lines.append(cell.text)

        # 页眉 / 页脚——每个 section 都可能有独立的首页 / 奇偶页页眉页脚，
        # 全部取出（带标记前缀，让 diff 里能看出是页眉还是正文变了）
        seen_in_body = {ln for ln in lines if ln}
        for sec in doc.sections:
            for label, part in (
                ("页眉", sec.header),
                ("首页页眉", sec.first_page_header),
                ("偶数页页眉", sec.even_page_header),
                ("页脚", sec.footer),
                ("首页页脚", sec.first_page_footer),
                ("偶数页页脚", sec.even_page_footer),
            ):
                try:
                    if getattr(part, "is_linked_to_previous", False):
                        continue  # 链接到上一节的页眉页脚没有独立内容，跳过避免重复
                    for p in part.paragraphs:
                        if p.text:
                            lines.append(f"[{label}] {p.text}")
                except Exception:
                    # 个别 docx 没有某类页眉页脚，python-docx 偶尔抛异常，跳过即可
                    continue

        # XML <w:t> 兜底——抓正文流之外的文本框 / 形状文字。
        # 只补"正文里没出现过"的新文本，避免把已展示的段落重复一遍。
        try:
            from docx.oxml.ns import qn

            extra: list[str] = []
            for t in doc.element.body.iter(qn("w:t")):
                txt = t.text
                if txt and txt not in seen_in_body:
                    extra.append(txt)
                    seen_in_body.add(txt)
            for txt in extra:
                lines.append(f"[文本框] {txt}")
        except Exception:
            # 抓 XML 文本节点是增强项，失败不影响主流程
            pass

        return "\n".join(lines)

    def describe_change(self, old_blob: bytes, new_blob: bytes) -> str:
        """
        返回 "增加 X 段 / 删除 Y 段 / 修改 Z 段"。
        # 人工注释：用 SequenceMatcher 的 opcodes 区分"修改"和"纯增/删"——
        # 'replace' 块按 max(old_len,new_len) 计入"修改"，
        # 'insert' / 'delete' 才记入"增加" / "删除"。
        """
        old_text = self.extract_text(old_blob)
        new_text = self.extract_text(new_blob)
        if old_text is None or new_text is None:
            return "[文件解析失败：无法读取 .docx 段落]"

        old_paras = old_text.splitlines()
        new_paras = new_text.splitlines()
        sm = difflib.SequenceMatcher(a=old_paras, b=new_paras, autojunk=False)
        added = removed = modified = 0
        for op, i1, i2, j1, j2 in sm.get_opcodes():
            if op == "insert":
                added += j2 - j1
            elif op == "delete":
                removed += i2 - i1
            elif op == "replace":
                modified += max(i2 - i1, j2 - j1)

        if added == 0 and removed == 0 and modified == 0:
            return "段落级一致（仅样式或非文本元素差异）"
        return f"增加 {added} 段 / 删除 {removed} 段 / 修改 {modified} 段"

    def render_diff(self, old_blob: bytes, new_blob: bytes) -> list[DiffLine]:
        """
        段落级 diff——每个段落 / 单元格一行。
        坏 .docx 返回单条 meta，让 UI 显式告诉用户"解析失败"。
        """
        if not old_blob and not new_blob:
            return []

        if not old_blob:
            new_text = self.extract_text(new_blob)
            if new_text is None:
                return [("meta", "[文件解析失败：无法读取 .docx 段落]")]
            return self._text_as_diff(new_text, "added")

        if not new_blob:
            old_text = self.extract_text(old_blob)
            if old_text is None:
                return [("meta", "[文件解析失败：无法读取 .docx 段落]")]
            return self._text_as_diff(old_text, "removed")

        old_text = self.extract_text(old_blob)
        new_text = self.extract_text(new_blob)
        if old_text is None or new_text is None:
            return [("meta", "[文件解析失败：无法读取 .docx 段落]")]

        old_paras = old_text.splitlines()
        new_paras = new_text.splitlines()
        result: list[DiffLine] = []
        for line in difflib.ndiff(old_paras, new_paras):
            tag2 = line[:2]
            content = line[2:]
            if tag2 == "  ":
                result.append(("normal", content))
            elif tag2 == "+ ":
                result.append(("added", content))
            elif tag2 == "- ":
                result.append(("removed", content))
            # '? ' 是 difflib 的位置提示，跳过

        return result


# 模块导入时自动注册
HandlerRegistry.register(DocxHandler())
